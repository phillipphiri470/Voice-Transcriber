"""
exporter.py — Export transcription results to TXT, DOCX, and SRT formats.
"""

import logging
from datetime import datetime
from pathlib import Path

log = logging.getLogger(__name__)


def _seconds_to_srt_time(seconds: float) -> str:
    """Convert float seconds to SRT timestamp format: HH:MM:SS,mmm"""
    h  = int(seconds // 3600)
    m  = int((seconds % 3600) // 60)
    s  = int(seconds % 60)
    ms = int((seconds - int(seconds)) * 1000)
    return f"{h:02d}:{m:02d}:{s:02d},{ms:03d}"


def export_txt(result: dict, output_base: str) -> str:
    """Export as plain text with optional speaker labels."""
    out_path = f"{output_base}.txt"
    lines = [
        "SPEECH TRANSCRIPTION",
        f"Generated : {datetime.now().strftime('%d %B %Y %H:%M')}",
        f"Duration  : {result['duration_seconds']:.1f} seconds",
        f"Language  : {result['language']}",
        "",
        "=" * 60,
        "",
    ]

    current_speaker = None
    for seg in result["segments"]:
        speaker = seg.get("speaker")
        if speaker and speaker != current_speaker:
            lines.append(f"\n[{speaker}]")
            current_speaker = speaker
        timestamp = f"[{_seconds_to_srt_time(seg['start'])[:8]}]"
        lines.append(f"{timestamp} {seg['text']}")

    Path(out_path).write_text("\n".join(lines), encoding="utf-8")
    return out_path


def export_srt(result: dict, output_base: str) -> str:
    """Export as an SRT subtitle file (usable in video editors and media players)."""
    out_path = f"{output_base}.srt"
    lines = []
    for i, seg in enumerate(result["segments"], start=1):
        start = _seconds_to_srt_time(seg["start"])
        end   = _seconds_to_srt_time(seg["end"])
        text  = seg["text"]
        if seg.get("speaker"):
            text = f"[{seg['speaker']}] {text}"
        lines.append(str(i))
        lines.append(f"{start} --> {end}")
        lines.append(text)
        lines.append("")

    Path(out_path).write_text("\n".join(lines), encoding="utf-8")
    return out_path


def export_docx(result: dict, output_base: str) -> str:
    """
    Export as a formatted Word document with a title page and timestamped paragraphs.
    Requires: pip install python-docx
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        log.error("python-docx not installed. pip install python-docx")
        return export_txt(result, output_base)

    doc = Document()

    # Title
    title = doc.add_heading("Speech Transcription", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata
    meta = doc.add_paragraph()
    meta.add_run(f"Generated: ").bold = True
    meta.add_run(datetime.now().strftime("%d %B %Y at %H:%M"))
    meta.add_run(f"\nDuration: ").bold = True
    meta.add_run(f"{result['duration_seconds']:.1f} seconds")
    meta.add_run(f"\nLanguage: ").bold = True
    meta.add_run(result["language"].upper())

    doc.add_paragraph()
    doc.add_heading("Transcript", level=1)

    current_speaker = None
    for seg in result["segments"]:
        speaker = seg.get("speaker")

        # New speaker heading
        if speaker and speaker != current_speaker:
            p = doc.add_heading(speaker, level=2)
            current_speaker = speaker

        # Timestamped paragraph
        p = doc.add_paragraph()
        timestamp_run = p.add_run(f"[{_seconds_to_srt_time(seg['start'])[:8]}] ")
        timestamp_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        timestamp_run.font.size = Pt(9)
        p.add_run(seg["text"])

    out_path = f"{output_base}.docx"
    doc.save(out_path)
    return out_path


def export_transcript(result: dict, output_base: str, fmt: str) -> str:
    """
    Dispatch export to the correct format handler.
    fmt: 'txt' | 'srt' | 'docx'
    """
    handlers = {
        "txt":  export_txt,
        "srt":  export_srt,
        "docx": export_docx,
    }
    handler = handlers.get(fmt)
    if handler is None:
        log.warning(f"Unknown export format '{fmt}'. Falling back to txt.")
        handler = export_txt
    return handler(result, output_base)
